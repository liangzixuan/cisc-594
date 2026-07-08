from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = "PantryPilot_Risk_Management_Report_Zixuan_Liang.docx"


COLORS = {
    "blue": "2E74B5",
    "dark_blue": "1F4D78",
    "header_fill": "F2F4F7",
    "light_fill": "F7F9FC",
    "border": "B8C2CC",
    "ink": "1F2933",
    "muted": "4B5563",
}


RISKS = [
    {
        "id": "R1",
        "category": "Safety / AI output",
        "title": "Unsafe LLM allergen output",
        "risk": "LLM-generated substitution or recipe includes an ingredient, alias, or derivative that violates a declared allergy.",
        "probability": 5,
        "impact": 5,
        "status": "Active",
        "priority": "Critical",
        "rationale": "The LLM is nondeterministic and user harm is possible if an unsafe ingredient bypasses the safety layer.",
        "mitigation": "Treat all LLM output as untrusted; apply deterministic allergen alias matching before display; block unsafe output; log near-misses; keep an allergen regression suite in CI.",
        "evidence": "REQ-PP-Q2-002, REQ-PP-Q3-003, allergen-violation regression tests, safety telemetry events.",
        "trigger": "Any blocked allergen output, alias-list miss, or recipe generated without safety-filter evidence.",
    },
    {
        "id": "R2",
        "category": "Algorithm / optimization",
        "title": "Optimizer infeasibility or timeout",
        "risk": "Nutrition optimizer cannot find a feasible weekly plan or consumes too much runtime under tight calorie, macro, allergy, and pantry constraints.",
        "probability": 4,
        "impact": 5,
        "status": "Active",
        "priority": "Critical",
        "rationale": "V2 combines multiple constraints; infeasible or slow solver behavior can create silent failure or unusable planning results.",
        "mitigation": "Bound solver runtime; define relaxation order; return partial plan with explicit relaxed-constraint notice; property-test generated constraint sets.",
        "evidence": "REQ-PP-Q1-001, REQ-PP-Q1-002, REQ-PP-Q3-002, feasibility tests and solver timeout checks.",
        "trigger": "Solver timeout, empty plan, conflicting target bands, or user-visible result without a relaxation explanation.",
    },
    {
        "id": "R3",
        "category": "External dependency",
        "title": "Recipe API instability",
        "risk": "Recipe API responses are incomplete, inconsistent, rate-limited, or temporarily unavailable, causing failed ingestion or unstable tests.",
        "probability": 4,
        "impact": 4,
        "status": "Active",
        "priority": "High",
        "rationale": "Spoonacular is outside project control and recipe fields may not match PantryPilot's canonical model.",
        "mitigation": "Normalize all recipe inputs; reject under-specified recipes; cache API responses; use recorded fixtures for tests; provide seed/manual recipes as fallbacks.",
        "evidence": "Recipe normalization pipeline, cached fixtures, parser tests, seed-library fallback.",
        "trigger": "API quota warning, missing ingredient quantity, missing nutrition values, or parser failure.",
    },
    {
        "id": "R4",
        "category": "Boundary / date logic",
        "title": "Pantry expiry boundary defect",
        "risk": "Pantry expiry logic mishandles exact dates, promoting expired items or failing to prioritize near-expiry items.",
        "probability": 3,
        "impact": 5,
        "status": "Active",
        "priority": "High",
        "rationale": "Date boundary defects directly affect food safety and can undermine the core waste-reduction feature.",
        "mitigation": "Define an inclusive three-day near-expiry window; prohibit expired items; test today, +3 days, and yesterday boundary cases.",
        "evidence": "REQ-PP-Q1-003, REQ-PP-Q2-003, REQ-PP-Q3-001, boundary-date system tests.",
        "trigger": "Any recipe suggestion using an expired pantry item or failing an exact-boundary date test.",
    },
    {
        "id": "R5",
        "category": "Privacy / data protection",
        "title": "Sensitive profile-data exposure",
        "risk": "Dietary profile, allergy, or health-target data is exposed through logs, configuration, test fixtures, or AI prompts.",
        "probability": 3,
        "impact": 4,
        "status": "Active",
        "priority": "High",
        "rationale": "The project stores sensitive preference and health-related data even though it is not a regulated medical device.",
        "mitigation": "Store secrets outside version control; redact logs; avoid real personal data in fixtures; review configuration files before baselines.",
        "evidence": "CM audit checklist, environment template, pull-request review for secrets and sensitive data.",
        "trigger": "Secret in repo, real personal data in tests, or prompt/log entry containing unnecessary sensitive profile details.",
    },
    {
        "id": "R6",
        "category": "Data correctness",
        "title": "Grocery aggregation error",
        "risk": "Unit conversion or grocery aggregation defects produce incorrect consolidated shopping lists.",
        "probability": 3,
        "impact": 3,
        "status": "Monitor",
        "priority": "Medium",
        "rationale": "Wrong quantities reduce user trust and can cause waste, but the impact is lower than allergen or expired-food failures.",
        "mitigation": "Use canonical units after ingestion; add integration tests for duplicate ingredients, mixed units, and rounding rules.",
        "evidence": "Normalization pipeline, grocery-list integration tests, release checklist for known conversion cases.",
        "trigger": "Mismatch between recipe ingredients and grocery list, or rounding behavior that changes planned servings materially.",
    },
    {
        "id": "R7",
        "category": "Scope / schedule",
        "title": "V2 scope expansion",
        "risk": "V2 AI generation, optimizer, and pantry-aware planning expand beyond the semester project capacity.",
        "probability": 2,
        "impact": 4,
        "status": "Retired / controlled",
        "priority": "Medium",
        "rationale": "The risk was significant early, but versioning and CM baselines separated V1 deliverables from higher-complexity V2 work.",
        "mitigation": "Keep V1 'Plan & Adapt' as the release baseline; open V2 work only after V1 tag, tests, and release notes are complete.",
        "evidence": "Project proposal V1/V2 boundary, CM baseline v1.0.0, release notes, branching strategy.",
        "trigger": "V2 feature merged before V1 baseline or required V1 acceptance test deferred by expansion work.",
    },
]


WEEKS = [
    {
        "week": "Week 1",
        "date": "Jun. 16",
        "focus": "Project proposal and initial scope",
        "updates": "Identified AI-safety, API-dependency, grocery-correctness, and scope risks. Set V1/V2 boundary to avoid mixing optimizer work into the first release.",
        "score_changes": "Opened R1=25, R3=16, R6=12, R7=16.",
    },
    {
        "week": "Week 2",
        "date": "Jun. 23",
        "focus": "Requirements and boundary analysis",
        "updates": "Added optimizer infeasibility and pantry-date risks after ABC requirements clarified calorie bands, macro bands, expiry windows, and allergen prohibitions.",
        "score_changes": "Opened R2=20 and R4=15. Reduced R6 from 12 to 9 after canonical-unit design.",
    },
    {
        "week": "Week 3",
        "date": "Jun. 30",
        "focus": "Risk scoring normalized",
        "updates": "Applied a 1-5 probability and impact scale across all project risks. Added privacy/data-protection risk because profiles include allergies and health targets.",
        "score_changes": "Opened R5=12. Kept R1=25 as top cross-version risk.",
    },
    {
        "week": "Week 4",
        "date": "Jul. 7",
        "focus": "Mitigation design",
        "updates": "Paired each major risk with a control: deterministic allergen filter, API fixture cache, bounded solver, expiry boundary tests, and configuration review.",
        "score_changes": "R3 remained 16 because dependency probability is still high; R7 reduced from 16 to 12 after version-scope controls.",
    },
    {
        "week": "Week 5",
        "date": "Jul. 14",
        "focus": "Configuration management package",
        "updates": "Established baseline v1.0.0, branch strategy, pull-request review, CI expectations, changelog, release notes, and configuration audit checklist.",
        "score_changes": "R7 reduced from 12 to 8 and marked controlled; R5 remained 12 pending regular secret/config review.",
    },
    {
        "week": "Week 6",
        "date": "Jul. 21",
        "focus": "System testing plan",
        "updates": "Mapped risks to tests: allergen regression, optimizer feasibility/property tests, API parser fixtures, pantry-date boundary tests, and grocery aggregation tests.",
        "score_changes": "R1, R2, R3, and R4 stayed active because implementation defects would still affect release readiness.",
    },
    {
        "week": "Week 7",
        "date": "Jul. 28",
        "focus": "Risk report baseline",
        "updates": "Reviewed the risk register for final reporting. No new top risk was added; R7 remained retired/controlled, while safety, optimizer, API, pantry-date, and privacy risks stayed active.",
        "score_changes": "Current tracked scores: R1=25, R2=20, R3=16, R4=15, R5=12, R6=9, R7=8.",
    },
]


TREND = {
    "R1": [25, 25, 25, 25, 25, 25, 25],
    "R2": ["-", 20, 20, 20, 20, 20, 20],
    "R3": [16, 16, 16, 16, 16, 16, 16],
    "R4": ["-", 15, 15, 15, 15, 15, 15],
    "R5": ["-", "-", 12, 12, 12, 12, 12],
    "R6": [12, 9, 9, 9, 9, 9, 9],
    "R7": [16, 16, 16, 12, 8, 8, 8],
}


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, width_dxa=9360, indent_dxa=120):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")


def set_column_widths(table, widths):
    for row in table.rows:
        for idx, width in enumerate(widths):
            cell = row.cells[idx]
            cell.width = Inches(width)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(int(width * 1440)))


def apply_table_style(table, widths=None, font_size=9):
    table.style = "Table Grid"
    table.autofit = False
    table.allow_autofit = False
    set_table_width(table)
    if widths:
        set_column_widths(table, widths)
    for r_idx, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.08
                for run in paragraph.runs:
                    run.font.name = "Calibri"
                    run.font.size = Pt(font_size)
                    run.font.color.rgb = RGBColor.from_string(COLORS["ink"])
            if r_idx == 0:
                set_cell_shading(cell, COLORS["header_fill"])
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True


def set_cell_text(cell, text, bold=False, size=9, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    if align:
        p.alignment = align
    run = p.add_run(str(text))
    run.bold = bold
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(COLORS["ink"])


def add_para(doc, text="", style=None, bold=False, italic=False):
    p = doc.add_paragraph(style=style)
    if text:
        run = p.add_run(text)
        run.bold = bold
        run.italic = italic
    return p


def add_section_title(doc, text):
    p = doc.add_paragraph(style="Heading 1")
    p.add_run(text)
    return p


def add_subtitle(doc, text):
    p = doc.add_paragraph(style="Heading 2")
    p.add_run(text)
    return p


def add_table_caption(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor.from_string(COLORS["dark_blue"])
    return p


def setup_styles(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(COLORS["ink"])
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    title = styles["Title"]
    title.font.name = "Calibri"
    title.font.size = Pt(22)
    title.font.bold = True
    title.font.color.rgb = RGBColor.from_string(COLORS["ink"])
    title.paragraph_format.space_after = Pt(4)

    for style_name, size, color, before, after in [
        ("Heading 1", 16, COLORS["blue"], 16, 8),
        ("Heading 2", 13, COLORS["blue"], 12, 6),
        ("Heading 3", 12, COLORS["dark_blue"], 8, 4),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def add_footer(doc):
    for section in doc.sections:
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer.add_run("PantryPilot Risk Management Report | Zixuan Liang")
        run.font.name = "Calibri"
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor.from_string(COLORS["muted"])


def build_report():
    doc = Document()
    setup_styles(doc)

    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("PantryPilot Risk Management Report")
    for text in [
        "Zixuan Liang",
        "Harrisburg University of Science and Technology",
        "CISC 594: Software Testing Principles and Techniques",
        "Dr. Khalid Lateef",
        "July 28, 2026",
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.size = Pt(12)
    add_para(doc)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Project: PantryPilot — Smart Meal and Grocery Planner")
    r.bold = True
    r.font.size = Pt(12)
    doc.add_page_break()

    add_section_title(doc, "Executive Summary")
    add_para(
        doc,
        "This report identifies, scores, tracks, and controls the major risks for PantryPilot, the semester project for CISC 594. PantryPilot is a smart meal and grocery planner that maintains dietary profiles, recipe data, pantry inventory, weekly plans, AI-assisted substitutions, nutrition targets, and grocery lists. The risk-management approach follows a probability-by-impact model so that safety, correctness, external-dependency, privacy, and schedule risks can be compared on one scale.",
    )
    add_para(
        doc,
        "The highest current risk is unsafe AI output because the system uses an LLM to propose substitutions and, in Version 2, generate recipes. The highest-priority control is therefore not another prompt; it is a deterministic safety layer that treats LLM output as untrusted and blocks allergen-containing suggestions before a user can see them. Other major risks are optimizer infeasibility, recipe API variability, pantry-expiry boundary defects, privacy exposure, grocery-list aggregation defects, and scope creep from V2 features.",
    )

    add_section_title(doc, "Risk Scoring Method")
    add_para(
        doc,
        "Each risk is assigned a probability score from 1 to 5 and an impact score from 1 to 5. The total risk score is calculated as probability multiplied by impact. This gives a maximum score of 25 and supports prioritization across risks that have different causes but compete for the same project time and testing effort.",
    )
    add_table_caption(doc, "Table 1")
    add_para(doc, "Risk score calculation and priority bands", italic=True)
    table = doc.add_table(rows=1, cols=4)
    headers = ["Score range", "Priority", "Interpretation", "Expected action"]
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, bold=True)
    rows = [
        ("20-25", "Critical", "Could block release or create serious safety/trust failure.", "Track every week; require design control and regression evidence."),
        ("12-19", "High", "Material project or user-impact risk.", "Mitigate, assign test evidence, and review before baseline."),
        ("6-11", "Medium", "Manageable but still relevant to quality or schedule.", "Monitor, test representative cases, and reduce when controls are effective."),
        ("1-5", "Low", "Unlikely or minor under current scope.", "Record if useful; retire when no longer relevant."),
    ]
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
    apply_table_style(table, [0.9, 0.9, 2.3, 2.2], font_size=9)

    add_section_title(doc, "Current Risk Identification and Assessment")
    add_para(
        doc,
        "The current register contains seven risks. Six remain active or monitored, while the scope/schedule risk is marked as controlled because the project adopted a V1/V2 release boundary and a CM baseline. R1 through R5 are the most important risks because they could affect user safety, system correctness, test stability, or privacy.",
    )
    add_table_caption(doc, "Table 2")
    add_para(doc, "Current PantryPilot risk register", italic=True)
    table = doc.add_table(rows=1, cols=7)
    headers = ["ID", "Category", "Risk focus", "P", "I", "Score", "Status"]
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER if i in [0, 3, 4, 5] else None)
    for risk in RISKS:
        cells = table.add_row().cells
        values = [
            risk["id"],
            risk["category"],
            risk["title"],
            risk["probability"],
            risk["impact"],
            risk["probability"] * risk["impact"],
            f"{risk['priority']}; {risk['status']}",
        ]
        for i, value in enumerate(values):
            set_cell_text(cells[i], value, align=WD_ALIGN_PARAGRAPH.CENTER if i in [0, 3, 4, 5] else None)
    apply_table_style(table, [0.42, 1.35, 2.05, 0.42, 0.42, 0.55, 1.12], font_size=8.8)

    add_section_title(doc, "Risk Rationale and Impact")
    for risk in RISKS:
        add_subtitle(doc, f"{risk['id']}: {risk['category']}")
        add_para(doc, f"Risk: {risk['risk']}")
        add_para(doc, f"Assessment: probability {risk['probability']} × impact {risk['impact']} = {risk['probability'] * risk['impact']} ({risk['priority']}). {risk['rationale']}")
        add_para(doc, f"Control trigger: {risk['trigger']}")

    doc.add_page_break()
    add_section_title(doc, "Risk Mitigation and Control")
    add_para(
        doc,
        "Mitigation is documented for each tracked risk. The control design favors verification evidence that can be repeated during configuration audits and release reviews: regression tests, boundary tests, fixture-based integration tests, CI checks, and release baselines.",
    )
    add_table_caption(doc, "Table 3")
    add_para(doc, "Mitigation plan and control evidence", italic=True)
    table = doc.add_table(rows=1, cols=4)
    headers = ["Risk", "Mitigation / control", "Verification evidence", "Owner / status"]
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, bold=True)
    for risk in RISKS:
        cells = table.add_row().cells
        values = [risk["id"], risk["mitigation"], risk["evidence"], f"Zixuan Liang; {risk['status']}"]
        for i, value in enumerate(values):
            set_cell_text(cells[i], value)
    apply_table_style(table, [0.52, 2.85, 2.05, 1.02], font_size=8.5)

    doc.add_page_break()
    add_section_title(doc, "Week-by-Week Risk Tracking")
    add_para(
        doc,
        "The report maintains historical risk records rather than only showing the final score. Each weekly review re-examined whether risks should be added, re-scored, mitigated, monitored, or retired. A dash in the trend table means the risk had not yet been opened in the register for that week.",
    )
    add_table_caption(doc, "Table 4")
    add_para(doc, "Weekly risk-management record", italic=True)
    table = doc.add_table(rows=1, cols=5)
    headers = ["Week", "Review date", "Project focus", "Risk update", "Score/action record"]
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, bold=True)
    for week in WEEKS:
        cells = table.add_row().cells
        values = [week["week"], week["date"], week["focus"], week["updates"], week["score_changes"]]
        for i, value in enumerate(values):
            set_cell_text(cells[i], value)
    apply_table_style(table, [0.58, 0.62, 1.15, 2.55, 1.55], font_size=8.1)

    add_table_caption(doc, "Table 5")
    add_para(doc, "Risk-score trend by week", italic=True)
    table = doc.add_table(rows=1, cols=9)
    headers = ["Risk", "W1", "W2", "W3", "W4", "W5", "W6", "W7", "Current status"]
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    status_lookup = {risk["id"]: risk["status"] for risk in RISKS}
    for risk_id, scores in TREND.items():
        cells = table.add_row().cells
        values = [risk_id, *scores, status_lookup[risk_id]]
        for i, value in enumerate(values):
            set_cell_text(cells[i], value, align=WD_ALIGN_PARAGRAPH.CENTER)
    apply_table_style(table, [0.55, 0.42, 0.42, 0.42, 0.42, 0.42, 0.42, 0.42, 2.45], font_size=8.2)

    add_section_title(doc, "Risk Retirement and Continuing Monitoring")
    add_para(
        doc,
        "R7 is the only retired or controlled risk in the current register. It was reduced because the project separated V1 from V2, created a configuration-management baseline, and defined release criteria before higher-risk V2 work begins. R6 is monitored because unit conversion and aggregation defects remain possible, but the mitigation is straightforward and the impact is lower than safety or privacy risks.",
    )
    add_para(
        doc,
        "R1, R2, R3, R4, and R5 remain active. These risks should be reviewed before any implementation baseline because they connect directly to safety filtering, optimizer behavior, external dependency stability, expiry-date boundaries, and sensitive user data. The project should not tag a release unless the relevant controls have current test or audit evidence.",
    )

    add_section_title(doc, "Conclusion")
    add_para(
        doc,
        "The most important lesson from the PantryPilot risk analysis is that the system's smart behavior must be bounded by disciplined SQA controls. LLM-generated recipes and substitutions can add value, but the safety case depends on deterministic filtering, traceable requirements, CM baselines, and repeatable system tests. The weekly record shows that risks were not treated as a one-time list; they were opened, rescored, mitigated, monitored, and retired as the project scope became clearer.",
    )

    add_section_title(doc, "References")
    refs = [
        "Liang, Z. (2026a). PantryPilot — A smart meal and grocery planner [Course project proposal]. Harrisburg University of Science and Technology.",
        "Liang, Z. (2026b). PantryPilot — Software requirements in ABC format [Course assignment]. Harrisburg University of Science and Technology.",
        "Liang, Z. (2026c). PantryPilot configuration management generic files [Course assignment]. Harrisburg University of Science and Technology.",
        "Liang, Z. (2026d). PantryPilot final project presentation [Course presentation]. Harrisburg University of Science and Technology.",
    ]
    for ref in refs:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.first_line_indent = Inches(-0.5)
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.space_after = Pt(6)
        p.add_run(ref)

    add_footer(doc)
    doc.save(OUT)


if __name__ == "__main__":
    build_report()
