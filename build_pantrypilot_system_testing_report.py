from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = "PantryPilot_System_Testing_Report_Zixuan_Liang_Updated.docx"

COLORS = {
    "blue": "2E74B5",
    "dark_blue": "1F4D78",
    "header_fill": "F2F4F7",
    "ink": "1F2933",
    "muted": "4B5563",
}


TEST_CASES = [
    {
        "id": "V1-ST-01",
        "version": "v1.0.0 / v1.1.0",
        "feature": "Release configuration",
        "procedure": "From the project root, load config/app_config.json and call describe_release(load_config()).",
        "expected": "The system displays the configured application name, version, and release name.",
        "actual": "PantryPilot 1.1.0: Safety & Grocery Controls",
        "result": "PASS",
    },
    {
        "id": "V1-ST-02",
        "version": "v1.0.0 / v1.1.0",
        "feature": "Quality gates",
        "procedure": "Open the runtime configuration and inspect quality_gates for unit_tests_required, pull_request_required, review_required, and ci_required.",
        "expected": "All four quality gates are enabled.",
        "actual": "ci_required, pull_request_required, review_required, unit_tests_required",
        "result": "PASS",
    },
    {
        "id": "V1-ST-03",
        "version": "v1.0.0 / v1.1.0",
        "feature": "Grocery list generation",
        "procedure": "Create two recipe inputs that both contain rice in cups, then call aggregate_grocery_items(recipes).",
        "expected": "Duplicate rice entries are normalized and summed to one grocery item with quantity 3.0 cup.",
        "actual": "[{'name': 'rice', 'quantity': 3.0, 'unit': 'cup'}]",
        "result": "PASS",
    },
    {
        "id": "V1.1-ST-01",
        "version": "v1.1.0",
        "feature": "Allergen safety filter",
        "procedure": "Create a generated recipe containing 'soy lecithin' and a user profile declaring soy allergy. Run detect_allergen_conflicts and recipe_is_safe_for_profile.",
        "expected": "The system detects soy and blocks the recipe as unsafe.",
        "actual": "{'conflicts': ['soy'], 'safe': False}",
        "result": "PASS",
    },
    {
        "id": "V1.1-ST-02",
        "version": "v1.1.0",
        "feature": "Safe recipe approval",
        "procedure": "Create a recipe with rice, broccoli, and olive oil for a profile with soy and peanut allergies. Run recipe_is_safe_for_profile.",
        "expected": "The system returns True because no declared allergen or alias appears in the recipe.",
        "actual": "True",
        "result": "PASS",
    },
    {
        "id": "V1.1-ST-03",
        "version": "v1.1.0",
        "feature": "Pantry expiry boundary",
        "procedure": "Evaluate is_near_expiry for today, +3 days, +4 days, and yesterday using Aug. 14, 2026 as the current date.",
        "expected": "Today and +3 days are included in the near-expiry window; yesterday and +4 days are excluded.",
        "actual": "{'today': True, 'plus_3': True, 'plus_4': False, 'yesterday': False}",
        "result": "PASS",
    },
]


WEB_TESTS = [
    ("WEB-01", "Demo page and release data", "GET / and GET /api/demo-data return the interface and controlled v1.1.0 release identity."),
    ("WEB-02", "Allergen decision endpoint", "POST /api/check-recipe blocks the tofu recipe when the household profile includes soy."),
    ("WEB-03", "Grocery aggregation endpoint", "POST /api/grocery-list consolidates shared ingredients and returns canonical totals."),
    ("WEB-04", "Pantry boundary endpoint", "POST /api/pantry-status applies the inclusive three-day window and identifies expired items."),
    ("WEB-05", "Quality evidence endpoint", "GET /api/quality-summary executes and reports all six release scenarios as passing."),
]


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, width_dxa=9360, indent_dxa=120):
    table.autofit = False
    table.allow_autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")


def set_column_widths(table, widths):
    for row in table.rows:
        for idx, width in enumerate(widths):
            cell = row.cells[idx]
            cell.width = Inches(width)
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(int(width * 1440)))


def set_cell_text(cell, text, bold=False, size=8.7, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    if align:
        p.alignment = align
    run = p.add_run(str(text))
    run.bold = bold
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(COLORS["ink"])


def apply_table_style(table, widths=None, font_size=8.7):
    table.style = "Table Grid"
    set_table_width(table)
    if widths:
        set_column_widths(table, widths)
    for row_index, row in enumerate(table.rows):
        tr_pr = row._tr.get_or_add_trPr()
        if tr_pr.find(qn("w:cantSplit")) is None:
            tr_pr.append(OxmlElement("w:cantSplit"))
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)
            if row_index == 0:
                set_cell_shading(cell, COLORS["header_fill"])
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.08
                for run in paragraph.runs:
                    run.font.name = "Calibri"
                    run.font.size = Pt(font_size)
                    run.font.color.rgb = RGBColor.from_string(COLORS["ink"])
                    if row_index == 0:
                        run.bold = True


def add_para(doc, text="", bold=False, italic=False):
    p = doc.add_paragraph()
    if text:
        r = p.add_run(text)
        r.bold = bold
        r.italic = italic
    return p


def add_heading(doc, text, level=1):
    return doc.add_paragraph(text, style=f"Heading {level}")


def add_caption(doc, label, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(label)
    r.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor.from_string(COLORS["dark_blue"])
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(4)
    r2 = p2.add_run(title)
    r2.italic = True


def setup_doc(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(COLORS["ink"])
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for style_name, size, color, before, after in [
        ("Heading 1", 16, COLORS["blue"], 16, 8),
        ("Heading 2", 13, COLORS["blue"], 12, 6),
        ("Heading 3", 12, COLORS["dark_blue"], 8, 4),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def add_footer(doc):
    for section in doc.sections:
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer.add_run("PantryPilot System Testing Report | Zixuan Liang")
        run.font.name = "Calibri"
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor.from_string(COLORS["muted"])


def add_title_page(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(92)
    run = p.add_run("PantryPilot System Testing Report")
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor.from_string(COLORS["ink"])

    rows = [
        ("Prepared by", "Zixuan Liang"),
        ("Course", "CISC 594: Software Testing Principles and Techniques"),
        ("Instructor", "Dr. Khalid Lateef"),
        ("Project", "PantryPilot - Smart Meal and Grocery Planner"),
        ("Submission", "System Testing Report"),
        ("Date", "August 14, 2026"),
    ]
    table = doc.add_table(rows=0, cols=2)
    for label, value in rows:
        cells = table.add_row().cells
        set_cell_text(cells[0], label, bold=True, size=10)
        set_cell_text(cells[1], value, size=10)
    apply_table_style(table, [1.55, 4.75], font_size=10)
    doc.add_page_break()


def add_test_overview_table(doc):
    add_caption(doc, "Table 3", "System test case overview")
    table = doc.add_table(rows=1, cols=5)
    headers = ["Test ID", "Version", "Feature", "System behavior tested", "Status"]
    for idx, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[idx], header, bold=True)
    for case in TEST_CASES:
        cells = table.add_row().cells
        values = [case["id"], case["version"], case["feature"], case["expected"], case["result"]]
        for idx, value in enumerate(values):
            set_cell_text(cells[idx], value, align=WD_ALIGN_PARAGRAPH.CENTER if idx in [0, 4] else None)
    apply_table_style(table, [0.78, 1.1, 1.3, 2.55, 0.65], font_size=8.5)


def add_web_test_table(doc):
    add_caption(doc, "Table 4", "Web presentation-layer regression tests")
    table = doc.add_table(rows=1, cols=4)
    headers = ["Test ID", "Workflow", "Expected system behavior", "Status"]
    for idx, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[idx], header, bold=True)
    for test_id, workflow, expected in WEB_TESTS:
        cells = table.add_row().cells
        values = [test_id, workflow, expected, "PASS"]
        for idx, value in enumerate(values):
            set_cell_text(cells[idx], value, align=WD_ALIGN_PARAGRAPH.CENTER if idx in [0, 3] else None)
    apply_table_style(table, [0.8, 1.55, 3.45, 0.65], font_size=8.4)


def add_test_case_detail_tables(doc, table_start=4):
    for offset, case in enumerate(TEST_CASES):
        add_caption(doc, f"Table {table_start + offset}", f"{case['id']} detailed system test procedure")
        table = doc.add_table(rows=0, cols=2)
        rows = [
            ("Version", case["version"]),
            ("Feature", case["feature"]),
            ("Preconditions", "Repository extracted; Python can import src.pantrypilot_app; configuration file is available. For v1.1 tests, safety and grocery functions are present on tag/main v1.1.0."),
            ("Test procedure", case["procedure"]),
            ("Expected result", case["expected"]),
            ("Actual result", case["actual"]),
            ("Status", case["result"]),
        ]
        for label, value in rows:
            cells = table.add_row().cells
            set_cell_text(cells[0], label, bold=True, size=8.8)
            set_cell_text(cells[1], value, size=8.8)
        apply_table_style(table, [1.25, 5.2], font_size=8.8)


def build_report():
    doc = Document()
    setup_doc(doc)
    add_title_page(doc)

    add_heading(doc, "Executive Summary")
    add_para(
        doc,
        "This report documents the system-level testing approach for PantryPilot across its current project versions and its web demonstration layer. The report includes the development and test environment, setup instructions, test-case selection methodology, detailed system test procedures, expected results, observed actual results, and final outcome. The tests exercise externally meaningful workflows including release configuration, quality gates, grocery-list generation, allergen blocking, safe recipe approval, pantry-expiry boundary handling, and browser-to-API integration.",
    )
    add_para(
        doc,
        "All six formal system scenarios passed, and all five web/API regression tests passed. Together, the pytest suite contains 11 passing automated tests. Local execution used Python 3.11.15, Flask 3.1.1, pytest 8.2.2, and Google Chrome 150.0.7871.129 on macOS 26.5.1. The configured CI environment uses GitHub Actions on ubuntu-latest with Python 3.12 and runs both pytest and the six-scenario system-test runner.",
    )

    add_heading(doc, "Versions Under Test")
    add_caption(doc, "Table 1", "Application versions and scope")
    table = doc.add_table(rows=1, cols=4)
    for idx, header in enumerate(["Version", "Release name", "Scope under test", "Release evidence"]):
        set_cell_text(table.rows[0].cells[idx], header, bold=True)
    rows = [
        ("v1.0.0", "Plan & Adapt", "Configuration baseline, feature flags, quality gates, and grocery-list generation behavior.", "Git tag v1.0.0; baseline commit 842ae5e."),
        ("v1.1.0", "Safety & Grocery Controls", "All v1.0.0 behavior plus allergen filtering, safe recipe approval, grocery aggregation, and pantry-expiry boundary logic.", "Git tag v1.1.0; merge commit 834e64e."),
        ("Unreleased", "Web demonstration layer", "Browser workflows and Flask API endpoints that expose the controlled v1.1.0 safety, grocery, pantry, and quality behavior.", "CHANGELOG [Unreleased]; tests/test_web_app.py."),
    ]
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], value)
    apply_table_style(table, [0.8, 1.45, 2.75, 1.45], font_size=8.7)

    doc.add_page_break()
    add_heading(doc, "Development and Test Environment")
    add_caption(doc, "Table 2", "Software and environment used")
    table = doc.add_table(rows=1, cols=3)
    for idx, header in enumerate(["Item", "Version / value", "Use in project"]):
        set_cell_text(table.rows[0].cells[idx], header, bold=True)
    rows = [
        ("Operating system", "macOS 26.5.1 arm64", "Local development and report test execution."),
        ("Local Python", "Python 3.11.15", "Executed system_test_runner.py for observed outputs."),
        ("CI Python", "Python 3.12", "Configured in GitHub Actions workflow."),
        ("Test framework", "pytest 8.2.2", "Pinned in requirements.txt and used by CI workflow."),
        ("Web framework", "Flask 3.1.1", "Serves the local demonstration UI and API endpoints."),
        ("Browser", "Google Chrome 150.0.7871.129", "Desktop and responsive presentation-layer verification."),
        ("Version control", "git version 2.50.1", "Controlled baselines, branches, and tags."),
        ("Application repository", "github.com/liangzixuan/cisc-594", "GitHub repository; PantryPilot is maintained in PantryPilot_CM_Generic_Files."),
        ("Current application config", "baseline_version 1.1.0; release_name Safety & Grocery Controls", "Runtime release identity."),
    ]
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], value)
    apply_table_style(table, [1.45, 2.1, 2.9], font_size=8.8)

    add_heading(doc, "Setup Instructions")
    add_para(doc, "A maintenance or grading reviewer can recreate the test environment using the following steps:")
    steps = [
        "Extract the submitted source-code package and open the PantryPilot_CM_Generic_Files directory.",
        "Confirm Git history with git log --oneline --decorate --graph --all and verify tags v1.0.0 and v1.1.0.",
        "Install the configured test dependency with python -m pip install -r requirements.txt if pytest is available.",
        "Run automated tests with python -m pytest. In CI, GitHub Actions performs this step on Python 3.12.",
        "Run the report evidence runner with python3 system_test_runner.py to reproduce the actual outputs listed in this report.",
        "Start the local web demonstration with python -m src.web_app, open http://127.0.0.1:5000, and exercise the Plan, Grocery, Pantry, and Quality views.",
        "For version-specific review, check out git tag v1.0.0 or v1.1.0 and repeat the relevant test procedures.",
    ]
    for step in steps:
        p = doc.add_paragraph(style="List Number")
        p.add_run(step)

    add_heading(doc, "Test-Case Selection Methodology")
    add_para(
        doc,
        "The formal system scenarios were selected from the highest-value user-visible behavior and the highest-risk requirements in the PantryPilot project. The goal was not to test every function in isolation; it was to confirm that each release behaves correctly from a system perspective. The six-scenario set therefore includes configuration loading, release identity, quality-gate configuration, grocery aggregation, allergen safety filtering, safe-output acceptance, and pantry expiry boundary behavior. Five additional Flask test-client cases verify that the demonstration interface reaches those same controls through stable HTTP endpoints.",
    )
    add_para(
        doc,
        "The selected cases provide adequate coverage for the current project scope because they exercise every major deterministic behavior in the submitted baseline and then verify its presentation through the web API. The allergen and expiry cases cover safety-critical negative and boundary behavior; the grocery case covers user-visible output correctness; configuration and quality-gate cases verify release readiness; web cases cover route availability, request/response integration, and evidence reporting. Manual browser checks at desktop and mobile widths supplement automation for layout and interaction quality. V2 nutrition optimization, external recipe ingestion, and generative-AI behavior remain future scope and are not represented as implemented features.",
    )

    doc.add_page_break()
    add_heading(doc, "System Test Plan and Results")
    add_test_overview_table(doc)
    doc.add_page_break()
    add_web_test_table(doc)
    add_test_case_detail_tables(doc, table_start=5)

    add_heading(doc, "Execution Summary")
    add_caption(doc, "Table 11", "Test execution outcome summary")
    table = doc.add_table(rows=1, cols=4)
    for idx, header in enumerate(["Version", "Tests executed", "Passed", "Errors / defects observed"]):
        set_cell_text(table.rows[0].cells[idx], header, bold=True)
    rows = [
        ("v1.0.0 / v1.1.0 shared baseline", "3", "3", "None observed."),
        ("v1.1.0 Safety & Grocery Controls", "3", "3", "None observed."),
        ("Unreleased web/API layer", "5", "5", "None observed."),
        ("Total automated pytest suite", "11", "11", "None observed."),
    ]
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], value, align=WD_ALIGN_PARAGRAPH.CENTER if idx in [1, 2] else None)
    apply_table_style(table, [2.2, 1.05, 0.85, 2.35], font_size=8.8)

    doc.add_page_break()
    add_heading(doc, "Limitations and Future System Tests")
    add_para(
        doc,
        "The current tests cover the implemented release baseline and the local demonstration layer; they do not claim that unavailable future features have been tested. The browser review was manual, while HTTP behavior was automated with Flask's test client. Nutrition optimization, live AI-generated recipes, external recipe API ingestion, authentication, persistence, and full pantry-aware planning remain future scope. Those features will require recorded provider responses, security and privacy tests, infeasible-target cases, multi-profile household tests, and full end-to-end deployment checks.",
    )

    add_heading(doc, "Conclusion")
    add_para(
        doc,
        "The PantryPilot system testing process demonstrates that the tagged application baseline and its web presentation layer can be installed, identified, configured, and exercised at the system level. Six of six formal scenarios and 11 of 11 automated pytest tests produced the expected outputs with no observed defects. The evidence covers the implemented release behavior, the browser-to-API path, and the most important user-facing and risk-driven controls without overstating unimplemented V2 scope.",
    )

    add_heading(doc, "References")
    refs = [
        "Liang, Z. (2026a). PantryPilot configuration management report [Course assignment]. Harrisburg University of Science and Technology.",
        "Liang, Z. (2026b). PantryPilot risk management report [Course assignment]. Harrisburg University of Science and Technology.",
        "Liang, Z. (2026c). PantryPilot source repository and system test runner [Course project files]. Harrisburg University of Science and Technology.",
    ]
    for ref in refs:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.5)
        p.paragraph_format.space_after = Pt(6)
        p.add_run(ref)

    add_footer(doc)
    doc.save(OUT)


if __name__ == "__main__":
    build_report()
