from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


SOURCE = Path("PantryPilot_Product_Requirements_Document_Zixuan_Liang.md")
OUTPUT = Path("PantryPilot_Product_Requirements_Document_Zixuan_Liang.docx")

COLORS = {
    "blue": "2E74B5",
    "dark_blue": "1F4D78",
    "ink": "1F2933",
    "muted": "5B6573",
    "light_fill": "F2F4F7",
    "callout_fill": "EAF2F8",
    "callout_border": "2E74B5",
    "implemented": "E8F3EC",
    "partial": "FFF4D6",
    "not_implemented": "FCE8E6",
    "white": "FFFFFF",
}


def set_run_font(run, *, name="Calibri", size=None, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color="D1D5DB", size="4"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:color"), color)


def remove_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "nil")


def set_table_geometry(table, widths_in, indent_dxa=120):
    widths_dxa = [int(width * 1440) for width in widths_in]
    total_dxa = sum(widths_dxa)
    table.autofit = False
    table.allow_autofit = False

    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(total_dxa))

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_ind.set(qn("w:w"), str(indent_dxa))

    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)

    for row in table.rows:
        tr_pr = row._tr.get_or_add_trPr()
        if tr_pr.find(qn("w:cantSplit")) is None:
            tr_pr.append(OxmlElement("w:cantSplit"))
        for index, cell in enumerate(row.cells):
            cell.width = Inches(widths_in[index])
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(widths_dxa[index]))


def repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = tr_pr.find(qn("w:tblHeader"))
    if tbl_header is None:
        tbl_header = OxmlElement("w:tblHeader")
        tr_pr.append(tbl_header)
    tbl_header.set(qn("w:val"), "true")


def add_page_number(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    value = OxmlElement("w:t")
    value.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, instruction, separate, value, end))
    set_run_font(run, size=9, color=COLORS["muted"])


def setup_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(COLORS["ink"])
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for style_name, size, color, before, after in (
        ("Heading 1", 16, COLORS["blue"], 16, 8),
        ("Heading 2", 13, COLORS["blue"], 12, 6),
        ("Heading 3", 12, COLORS["dark_blue"], 8, 4),
    ):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    header.paragraph_format.space_after = Pt(0)
    header.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT)
    left = header.add_run("PANTRYPILOT")
    set_run_font(left, size=8.5, color=COLORS["muted"], bold=True)
    right = header.add_run("\tPRODUCT REQUIREMENTS DOCUMENT")
    set_run_font(right, size=8.5, color=COLORS["muted"])

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.paragraph_format.space_before = Pt(0)
    footer.paragraph_format.space_after = Pt(0)
    prefix = footer.add_run("Zixuan Liang | CISC 594 | Page ")
    set_run_font(prefix, size=9, color=COLORS["muted"])
    add_page_number(footer)


def add_inline_runs(paragraph, text, *, size=None, color=None):
    pattern = re.compile(r"(\*\*.*?\*\*|`.*?`)")
    position = 0
    for match in pattern.finditer(text):
        if match.start() > position:
            run = paragraph.add_run(text[position:match.start()])
            set_run_font(run, size=size, color=color or COLORS["ink"])
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, size=size, color=color or COLORS["ink"], bold=True)
        else:
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, name="Courier New", size=(size or 10) - 0.2, color=color or COLORS["ink"])
        position = match.end()
    if position < len(text):
        run = paragraph.add_run(text[position:])
        set_run_font(run, size=size, color=color or COLORS["ink"])


def add_masthead(doc):
    kicker = doc.add_paragraph()
    kicker.paragraph_format.space_before = Pt(10)
    kicker.paragraph_format.space_after = Pt(4)
    run = kicker.add_run("PRODUCT REQUIREMENTS DOCUMENT")
    set_run_font(run, size=10, color=COLORS["blue"], bold=True)

    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(5)
    run = title.add_run("PantryPilot")
    set_run_font(run, size=28, color=COLORS["ink"], bold=True)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_before = Pt(0)
    subtitle.paragraph_format.space_after = Pt(16)
    run = subtitle.add_run("Code-grounded as-built specification for the safety-aware meal and grocery planning demonstration")
    set_run_font(run, size=13, color=COLORS["muted"])

    metadata = [
        ("Document status", "As-built specification"),
        ("Backend baseline", "v1.1.0 - Safety & Grocery Controls"),
        ("Web layer", "Unreleased"),
        ("Owner", "Zixuan Liang"),
        ("Version and date", "PRD 1.0 | July 21, 2026"),
        ("Source", "PantryPilot_CM_Generic_Files code, tests, configuration, and release records"),
    ]
    table = doc.add_table(rows=0, cols=2)
    for label, value in metadata:
        cells = table.add_row().cells
        p = cells[0].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(label)
        set_run_font(run, size=9.5, color=COLORS["dark_blue"], bold=True)
        p = cells[1].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        add_inline_runs(p, value, size=9.5)
        for cell in cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell, top=45, bottom=45)
    set_table_geometry(table, [1.45, 5.05], indent_dxa=0)
    remove_table_borders(table)

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(6)


def add_callout(doc, text):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.18)
    paragraph.paragraph_format.right_indent = Inches(0.18)
    paragraph.paragraph_format.space_before = Pt(5)
    paragraph.paragraph_format.space_after = Pt(10)
    paragraph.paragraph_format.line_spacing = 1.10
    p_pr = paragraph._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), COLORS["callout_fill"])
    p_pr.append(shading)
    borders = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), COLORS["callout_border"])
    borders.append(left)
    p_pr.append(borders)
    add_inline_runs(paragraph, text, size=11, color=COLORS["dark_blue"])


def next_numbering_ids(doc):
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(node.get(qn("w:abstractNumId"))) for node in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    return (max(abstract_ids, default=-1) + 1, max(num_ids, default=0) + 1)


def create_abstract_numbering(doc, *, numbered):
    numbering = doc.part.numbering_part.element
    abstract_id, _ = next_numbering_ids(doc)
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "decimal" if numbered else "bullet")
    level.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "%1." if numbered else "•")
    level.append(lvl_text)
    lvl_jc = OxmlElement("w:lvlJc")
    lvl_jc.set(qn("w:val"), "left")
    level.append(lvl_jc)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "720")
    tabs.append(tab)
    p_pr.append(tabs)
    indent = OxmlElement("w:ind")
    indent.set(qn("w:left"), "720")
    indent.set(qn("w:hanging"), "360")
    p_pr.append(indent)
    level.append(p_pr)
    abstract.append(level)
    numbering.append(abstract)
    return abstract_id


def create_num_instance(doc, abstract_id):
    numbering = doc.part.numbering_part.element
    _, num_id = next_numbering_ids(doc)
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    level_override = OxmlElement("w:lvlOverride")
    level_override.set(qn("w:ilvl"), "0")
    start_override = OxmlElement("w:startOverride")
    start_override.set(qn("w:val"), "1")
    level_override.append(start_override)
    num.append(level_override)
    numbering.append(num)
    return num_id


def add_list_item(doc, text, num_id, *, compact=False):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    # Named override: sequential workflows use tighter spacing to stay together.
    paragraph.paragraph_format.space_after = Pt(5 if compact else 8)
    paragraph.paragraph_format.line_spacing = 1.167
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.extend((ilvl, num))
    p_pr.append(num_pr)
    add_inline_runs(paragraph, text)


def clean_table_text(value):
    return value.replace("**", "").replace("`", "").strip()


def table_layout(headers):
    first = headers[0].lower()
    if first == "capability":
        return [1.35, 1.15, 4.0], 8.5, "Current capability scope"
    if first == "id" and len(headers) == 3:
        return [0.72, 2.48, 3.30], 8.25, "Functional requirements and acceptance criteria"
    if first == "method and path":
        return [1.12, 1.12, 1.18, 1.65, 1.43], 7.55, "HTTP API contracts"
    if first == "evidence":
        return [1.75, 1.85, 2.90], 8.2, "Requirements-to-evidence traceability"
    if first == "id" and headers[0].startswith("ID"):
        return [0.72, 2.48, 3.30], 8.25, "Requirements"
    return [6.5 / len(headers)] * len(headers), 8.4, "Product specification matrix"


def add_table_caption(doc, number, title):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run(f"Table {number}. ")
    set_run_font(run, size=9.5, color=COLORS["dark_blue"], bold=True)
    run = paragraph.add_run(title)
    set_run_font(run, size=9.5, color=COLORS["muted"], italic=True)


def add_markdown_table(doc, rows, table_number, section_title):
    headers = [clean_table_text(cell) for cell in rows[0]]
    widths, font_size, default_caption = table_layout(headers)
    if headers[0] == "ID":
        if section_title == "Nonfunctional Requirements":
            caption = "Nonfunctional requirements and current evidence"
        else:
            caption = f"{section_title}: requirements and acceptance criteria"
    else:
        caption = default_caption
    add_table_caption(doc, table_number, caption)

    table = doc.add_table(rows=1, cols=len(headers))
    for index, header in enumerate(headers):
        paragraph = table.rows[0].cells[index].paragraphs[0]
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(header)
        set_run_font(run, size=font_size, color=COLORS["ink"], bold=True)
        set_cell_shading(table.rows[0].cells[index], COLORS["light_fill"])
    repeat_table_header(table.rows[0])

    for source_row in rows[2:]:
        if len(source_row) != len(headers):
            continue
        cells = table.add_row().cells
        for index, value in enumerate(source_row):
            text = clean_table_text(value)
            paragraph = cells[index].paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.04
            if index == 0 and (text.startswith("FR-") or text.startswith("NFR-")):
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_inline_runs(paragraph, text, size=font_size)
            cells[index].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cells[index])
            if "Current status" in headers and index == 1:
                if text.startswith("Implemented"):
                    set_cell_shading(cells[index], COLORS["implemented"])
                elif text == "Partial":
                    set_cell_shading(cells[index], COLORS["partial"])
                elif text.startswith("Not implemented"):
                    set_cell_shading(cells[index], COLORS["not_implemented"])
    for cell in table.rows[0].cells:
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_margins(cell)
    set_table_geometry(table, widths)
    set_table_borders(table)

    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(2)


def parse_table(lines, start):
    table_lines = []
    index = start
    while index < len(lines) and lines[index].strip().startswith("|"):
        raw = lines[index].strip().strip("|")
        table_lines.append([cell.strip() for cell in raw.split("|")])
        index += 1
    return table_lines, index


def add_heading(doc, text, level):
    paragraph = doc.add_paragraph(style=f"Heading {level}")
    add_inline_runs(paragraph, text, color=COLORS["blue"] if level < 3 else COLORS["dark_blue"])
    return paragraph


def add_body(doc, text):
    paragraph = doc.add_paragraph()
    add_inline_runs(paragraph, text)
    return paragraph


def build_document():
    markdown = SOURCE.read_text(encoding="utf-8")
    lines = markdown.splitlines()

    doc = Document()
    setup_document(doc)
    add_masthead(doc)

    bullet_abstract = create_abstract_numbering(doc, numbered=False)
    number_abstract = create_abstract_numbering(doc, numbered=True)
    active_list_type = None
    active_num_id = None
    started = False
    section_title = "Product requirements"
    table_number = 0

    index = 0
    while index < len(lines):
        raw = lines[index]
        stripped = raw.strip()

        if stripped.startswith("## 1. Purpose"):
            started = True
        if not started:
            index += 1
            continue

        if not stripped:
            active_list_type = None
            active_num_id = None
            index += 1
            continue

        if stripped.startswith("|"):
            rows, index = parse_table(lines, index)
            table_number += 1
            add_markdown_table(doc, rows, table_number, section_title)
            active_list_type = None
            active_num_id = None
            continue

        if stripped.startswith("```"):
            code_lines = []
            index += 1
            while index < len(lines) and not lines[index].strip().startswith("```"):
                code_lines.append(lines[index])
                index += 1
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.left_indent = Inches(0.25)
            paragraph.paragraph_format.space_before = Pt(4)
            paragraph.paragraph_format.space_after = Pt(8)
            for line_number, code_line in enumerate(code_lines):
                run = paragraph.add_run(code_line)
                set_run_font(run, name="Courier New", size=9.5, color=COLORS["ink"])
                if line_number < len(code_lines) - 1:
                    run.add_break(WD_BREAK.LINE)
            index += 1
            active_list_type = None
            active_num_id = None
            continue

        heading_match = re.match(r"^(#{2,4})\s+(.+)$", stripped)
        if heading_match:
            hashes, text = heading_match.groups()
            if text.startswith("Appendix A"):
                doc.add_page_break()
            level = len(hashes) - 1
            section_title = re.sub(r"^\d+(?:\.\d+)*\.\s*", "", text)
            add_heading(doc, text, level)
            active_list_type = None
            active_num_id = None
            index += 1
            continue

        if stripped.startswith(">"):
            quote_lines = []
            while index < len(lines) and (lines[index].strip().startswith(">") or not lines[index].strip()):
                if lines[index].strip().startswith(">"):
                    quote_lines.append(lines[index].strip()[1:].strip())
                index += 1
            add_callout(doc, " ".join(quote_lines))
            active_list_type = None
            active_num_id = None
            continue

        bullet_match = re.match(r"^-\s+(.+)$", stripped)
        number_match = re.match(r"^\d+\.\s+(.+)$", stripped)
        if bullet_match or number_match:
            list_type = "bullet" if bullet_match else "number"
            text = (bullet_match or number_match).group(1)
            if active_list_type != list_type or active_num_id is None:
                abstract_id = bullet_abstract if list_type == "bullet" else number_abstract
                active_num_id = create_num_instance(doc, abstract_id)
                active_list_type = list_type
            add_list_item(doc, text, active_num_id, compact=list_type == "number")
            index += 1
            continue

        add_body(doc, stripped)
        active_list_type = None
        active_num_id = None
        index += 1

    doc.core_properties.title = "PantryPilot Product Requirements Document"
    doc.core_properties.subject = "Code-grounded as-built product requirements"
    doc.core_properties.author = "Zixuan Liang"
    doc.core_properties.keywords = "PantryPilot, PRD, requirements, SQA, Flask"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_document()
