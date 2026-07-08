from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path("PantryPilot_CM_Submission_Zixuan_Liang.docx")


def add_page_number(section):
    paragraph = section.header.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(fld_end)


def setup_styles(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    add_page_number(section)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.line_spacing = 1.10
    normal.paragraph_format.space_after = Pt(6)

    for name, size, color in [
        ("Heading 1", 16, "2E74B5"),
        ("Heading 2", 13, "2E74B5"),
        ("Heading 3", 12, "1F4D78"),
    ]:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(12 if name != "Heading 3" else 8)
        style.paragraph_format.space_after = Pt(6 if name != "Heading 3" else 4)
        style.paragraph_format.line_spacing = 1.10


def para(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.line_spacing = 1.10
    p.paragraph_format.space_after = Pt(6)
    return p


def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.line_spacing = 1.10
    p.paragraph_format.space_after = Pt(6)
    p.add_run(text)
    return p


def cell_text(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.line_spacing = 1.05
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Calibri"
    run.font.size = Pt(10)


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        cell_text(table.rows[0].cells[idx], header, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for idx, item in enumerate(row):
            cell_text(cells[idx], item)
    doc.add_paragraph()


def build():
    doc = Document()
    setup_styles(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("PantryPilot Configuration Management Exercise")
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor.from_string("0B2545")

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("Zixuan Liang | CISC 594 | July 14, 2026")

    doc.add_heading("Submission Overview", level=1)
    para(
        doc,
        "This submission applies the configuration-management example document to the PantryPilot semester project using generic project files. The package demonstrates repository setup, controlled configuration items, branching and pull-request workflow, automated CI checks, and release baseline management.",
    )

    doc.add_heading("Project Context", level=1)
    para(
        doc,
        "PantryPilot is a smart meal and grocery planner. Its controlled artifacts include source code, runtime configuration, project documentation, tests, CI workflow files, changelog entries, and release notes. The generic files in this submission are intentionally lightweight, but they show how the real project would be managed under Git/GitHub configuration management.",
    )

    doc.add_heading("CM Levels Demonstrated", level=1)
    add_table(
        doc,
        ["Level", "What the example required", "PantryPilot artifact"],
        [
            (
                "Beginner",
                "Create a repository, add generic files, commit, and push.",
                "README, VERSION, config/app_config.json, src/pantrypilot_app.py, tests/test_smoke.py.",
            ),
            (
                "Intermediate",
                "Use branches, pull requests, review, and conflict-control practices.",
                "docs/branching_strategy.md, .github/pull_request_template.md, docs/git_commands_used.md.",
            ),
            (
                "Expert",
                "Use CI/CD workflow, tags, releases, and automated checks.",
                ".github/workflows/ci.yml, CHANGELOG.md, releases/release_notes_v1.0.0.md.",
            ),
        ],
    )

    doc.add_heading("Configuration Items", level=1)
    add_table(
        doc,
        ["Configuration item", "File or folder", "CM purpose"],
        [
            ("Source baseline", "src/", "Keeps application behavior under version control."),
            ("Tests", "tests/", "Provides repeatable verification for changes."),
            ("Runtime config", "config/", "Separates controlled settings from code."),
            ("CM documentation", "docs/", "Defines branching, change-control, and audit rules."),
            ("Automation", ".github/workflows/ci.yml", "Runs tests for pushes and pull requests."),
            ("Release records", "VERSION, CHANGELOG.md, releases/", "Tracks approved baselines and release history."),
        ],
    )

    doc.add_heading("Change-Control Workflow", level=1)
    bullet(doc, "Create a feature or fix branch from main.")
    bullet(doc, "Modify only the configuration items needed for the change.")
    bullet(doc, "Run tests locally before opening a pull request.")
    bullet(doc, "Use the pull-request template to document change purpose, affected files, test evidence, and rollback notes.")
    bullet(doc, "Merge only after review and passing CI checks.")
    bullet(doc, "Update changelog and tag a release when a new baseline is approved.")

    doc.add_heading("Verification", level=1)
    para(
        doc,
        "The generic source/config path was smoke-checked locally by loading the PantryPilot configuration and verifying the baseline release description. The included GitHub Actions workflow installs requirements and runs pytest in a GitHub-hosted environment.",
    )

    doc.add_heading("Submission Contents", level=1)
    bullet(doc, "PantryPilot_CM_Generic_Files.zip: complete generic CM package for upload.")
    bullet(doc, "PantryPilot_CM_Generic_Files/: unpacked source folder for inspection.")
    bullet(doc, "PantryPilot_CM_Submission_Zixuan_Liang.pdf: summary document.")

    doc.save(OUT)


if __name__ == "__main__":
    build()
