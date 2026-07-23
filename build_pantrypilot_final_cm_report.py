from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = "PantryPilot_Final_Configuration_Management_Report_Zixuan_Liang_Updated.docx"

COLORS = {
    "blue": "2E74B5",
    "dark_blue": "1F4D78",
    "header_fill": "F2F4F7",
    "ink": "1F2933",
    "muted": "4B5563",
}


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, width_dxa=9360, indent_dxa=120):
    table.autofit = False
    table.allow_autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")


def set_column_widths(table, widths):
    for row in table.rows:
        for idx, width in enumerate(widths):
            cell = row.cells[idx]
            cell.width = Inches(width)
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(int(width * 1440)))


def set_cell_text(cell, text, bold=False, size=9, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    if align:
        p.alignment = align
    run = p.add_run(str(text))
    run.bold = bold
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(COLORS["ink"])


def apply_table_style(table, widths=None, font_size=9):
    table.style = "Table Grid"
    set_table_width(table)
    if widths:
        set_column_widths(table, widths)
    for row_index, row in enumerate(table.rows):
        tr_pr = row._tr.get_or_add_trPr()
        if tr_pr.find(qn("w:cantSplit")) is None:
            tr_pr.append(OxmlElement("w:cantSplit"))
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)
            if row_index == 0:
                set_cell_shading(cell, COLORS["header_fill"])
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.08
                for run in paragraph.runs:
                    run.font.name = "Calibri"
                    run.font.size = Pt(font_size)
                    run.font.color.rgb = RGBColor.from_string(COLORS["ink"])
                    if row_index == 0:
                        run.bold = True


def add_para(doc, text="", style=None, bold=False, italic=False):
    p = doc.add_paragraph(style=style)
    if text:
        r = p.add_run(text)
        r.bold = bold
        r.italic = italic
    return p


def add_heading(doc, text, level=1):
    return doc.add_paragraph(text, style=f"Heading {level}")


def add_caption(doc, label, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(label)
    r.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor.from_string(COLORS["dark_blue"])
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(4)
    r2 = p2.add_run(title)
    r2.italic = True
    return p2


def setup_doc(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(COLORS["ink"])
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for style_name, size, color, before, after in [
        ("Heading 1", 16, COLORS["blue"], 16, 8),
        ("Heading 2", 13, COLORS["blue"], 12, 6),
        ("Heading 3", 12, COLORS["dark_blue"], 8, 4),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def add_footer(doc):
    for section in doc.sections:
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer.add_run("PantryPilot Configuration Management Report | Zixuan Liang")
        run.font.name = "Calibri"
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor.from_string(COLORS["muted"])


def add_title_page(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(88)
    run = p.add_run("PantryPilot Configuration Management Report")
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor.from_string(COLORS["ink"])

    rule = doc.add_paragraph()
    rule.paragraph_format.space_after = Pt(12)
    p_pr = rule._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), COLORS["blue"])
    p_bdr.append(bottom)
    p_pr.append(p_bdr)

    rows = [
        ("Prepared by", "Zixuan Liang"),
        ("Course", "CISC 594: Software Testing Principles and Techniques"),
        ("Instructor", "Dr. Khalid Lateef"),
        ("Project", "PantryPilot - Smart Meal and Grocery Planner"),
        ("Submission", "HU14 Final Project: Configuration Management Report"),
        ("Date", "August 4, 2026"),
    ]
    table = doc.add_table(rows=0, cols=2)
    for label, value in rows:
        cells = table.add_row().cells
        set_cell_text(cells[0], label, bold=True, size=10)
        set_cell_text(cells[1], value, size=10)
    apply_table_style(table, [1.5, 4.8], font_size=10)
    doc.add_page_break()


def build_report():
    doc = Document()
    setup_doc(doc)
    add_title_page(doc)

    add_heading(doc, "Executive Summary")
    add_para(
        doc,
        "This report documents how configuration management was applied to PantryPilot, the semester project for CISC 594. The complete CISC-594 workspace is maintained in a GitHub repository, while the PantryPilot application remains a clearly bounded configuration-item directory. Approved baselines use main, controlled feature development, a test-before-merge workflow, and annotated release tags. A source-code ZIP may accompany the report for LMS submission, while the GitHub history provides the primary audit trail.",
    )
    add_para(
        doc,
        "PantryPilot has two traceable baselines: v1.0.0 for the initial Plan & Adapt CM baseline and v1.1.0 for the controlled Safety & Grocery Controls update. The v1.1.0 update was developed on feature/safety-grocery-controls, tested, merged to main, and tagged after verification. The new Flask demonstration interface is documented in CHANGELOG.md under [Unreleased]; it exposes the v1.1.0 controls without changing the tagged baseline and must complete the same branch, test, review, and release process before a future tag.",
    )

    add_heading(doc, "Project Context and CM Objectives")
    add_para(
        doc,
        "PantryPilot is a smart meal and grocery planner that maintains dietary constraints, recipe data, pantry inventory, weekly plans, grocery lists, and AI-assisted substitutions. The local web application adds a presentation layer for demonstrating these controls through Plan, Grocery, Pantry, and Quality views. Because the project handles safety-sensitive constraints such as allergens and expiry dates, configuration management is needed to keep backend code, web assets, requirements, tests, configuration files, and release records synchronized.",
    )
    add_para(
        doc,
        "The CM objective was to make every change traceable from a reason for change to branch work, test evidence, merge, release tag, and rollback point. This matches the assignment requirement to use version control, develop outside the main branch, test changes before merging, and tag completed software versions.",
    )

    doc.add_page_break()
    add_heading(doc, "Repository and Tool Use")
    add_caption(doc, "Table 1", "Version-control evidence")
    table = doc.add_table(rows=1, cols=3)
    for idx, header in enumerate(["CM item", "Evidence", "Purpose"]):
        set_cell_text(table.rows[0].cells[idx], header, bold=True)
    rows = [
        ("Version-control tool", "GitHub repository github.com/liangzixuan/cisc-594; PantryPilot files are in PantryPilot_CM_Generic_Files.", "Provides auditable history, branch control, tags, and rollback."),
        ("Main branch", "main", "Approved working baseline after tested changes are merged."),
        ("Development branch", "feature/safety-grocery-controls", "Isolated new development from the baseline until verification."),
        ("Release tags", "v1.0.0 and v1.1.0", "Identify tested software versions and rollback points."),
        ("CI workflow", ".github/workflows/pantrypilot-ci.yml", "Runs pytest and the system-test runner for PantryPilot changes."),
        ("Instructor access", "GitHub repository plus an LMS source ZIP when required.", "Allows review of source, tests, branches, tags, workflow, and release records."),
    ]
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
    apply_table_style(table, [1.55, 2.65, 2.25], font_size=8.8)

    doc.add_page_break()
    add_heading(doc, "Configuration Items")
    add_para(
        doc,
        "The project treats backend and web source files, tests, static assets, runtime configuration, CM documentation, CI workflow, release notes, and version files as controlled configuration items. These items are tracked together because a release is only meaningful if the code, user interface, tests, configuration, and release description agree.",
    )
    add_caption(doc, "Table 2", "Controlled configuration items")
    table = doc.add_table(rows=1, cols=4)
    for idx, header in enumerate(["Configuration item", "Location", "Control method", "Release relevance"]):
        set_cell_text(table.rows[0].cells[idx], header, bold=True)
    rows = [
        ("Backend source", "src/pantrypilot_app.py", "Branch, commit, review, test", "Implements deterministic release behavior."),
        ("Web application", "src/web_app.py; templates/; static/", "Branch, review, API tests, browser review", "Presents release controls through the demo UI."),
        ("Tests", "tests/test_smoke.py; tests/test_web_app.py", "Required before merge", "Verifies backend and HTTP integration behavior."),
        ("Runtime configuration", "config/app_config.json", "Version controlled", "States baseline version and enabled features."),
        ("Release identity", "VERSION, CHANGELOG.md", "Updated with each release", "Keeps version status accounting current."),
        ("Release notes", "releases/", "One file per tagged release", "Documents scope, verification, and rollback."),
        ("Dependencies", "requirements.txt", "Pinned and reviewed", "Controls Flask 3.1.1 and pytest 8.2.2."),
        ("CI workflow", ".github/workflows/pantrypilot-ci.yml", "Protected repository file", "Runs repeatable pytest and system checks."),
        ("CM documentation", "docs/", "Reviewed before baseline", "Explains branching and change-control policy."),
    ]
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
    apply_table_style(table, [1.35, 1.65, 1.55, 1.9], font_size=8.5)

    doc.add_page_break()
    add_heading(doc, "Formal Change-Control Process")
    add_para(
        doc,
        "The formal change-control process is intentionally lightweight but complete enough for the semester project. A change is not merged just because it compiles; it must have a stated purpose, a branch, test evidence, updated configuration records, and a rollback path.",
    )
    add_caption(doc, "Table 3", "Change-control workflow")
    table = doc.add_table(rows=1, cols=4)
    for idx, header in enumerate(["Step", "Activity", "PantryPilot evidence", "Control value"]):
        set_cell_text(table.rows[0].cells[idx], header, bold=True)
    rows = [
        ("1", "Identify change request", "Safety controls were required for v1.1.0; a local web demonstration was later requested for final presentation evidence.", "Prevents unmanaged scope drift."),
        ("2", "Create branch from main", "feature/safety-grocery-controls", "Keeps new development separate from baseline."),
        ("3", "Implement and update configuration items", "Backend source, Flask routes, template, JavaScript, CSS, local assets, tests, requirements, and changelog are controlled together.", "Keeps executable behavior and records synchronized."),
        ("4", "Test change set", "11 pytest tests and 6 system scenarios passed locally; CI runs both commands.", "Verifies backend and presentation-layer behavior before merge."),
        ("5", "Merge to main", "Merge commit 834e64e", "Moves only tested change into approved baseline."),
        ("6", "Tag release", "Annotated tag v1.1.0", "Creates a recoverable version reference."),
        ("7", "Record status accounting", "CHANGELOG.md and releases/release_notes_v1.1.0.md", "Documents what changed and how to roll back."),
        ("8", "Hold unreleased work", "The Flask demo remains in CHANGELOG [Unreleased] until branch, review, merge, and tag criteria are complete.", "Avoids silently redefining the v1.1.0 tag."),
    ]
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
    apply_table_style(table, [0.45, 1.65, 2.65, 1.65], font_size=8.2)

    doc.add_page_break()
    add_heading(doc, "Branching, Merge, and Tagging Evidence")
    add_para(
        doc,
        "The repository history demonstrates the required branch workflow. The initial baseline was committed on main and tagged v1.0.0. New code was then developed on feature/safety-grocery-controls, committed as a branch change, merged back to main, and tagged v1.1.0 after verification.",
    )
    add_caption(doc, "Table 4", "Repository history")
    table = doc.add_table(rows=1, cols=4)
    for idx, header in enumerate(["Commit / tag", "Branch", "Description", "CM meaning"]):
        set_cell_text(table.rows[0].cells[idx], header, bold=True)
    rows = [
        ("842ae5e; tag v1.0.0", "main", "Establish PantryPilot CM baseline", "Initial approved baseline."),
        ("72a997e", "feature/safety-grocery-controls", "Add safety and grocery control tests", "Controlled branch implementation."),
        ("834e64e; tag v1.1.0", "main", "Merge safety and grocery controls", "Merged tested release baseline."),
        ("51e7a65", "main", "Preserve PantryPilot project history in the full CISC-594 repository", "Retains earlier branch and tag evidence after repository consolidation."),
    ]
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
    apply_table_style(table, [1.25, 1.65, 1.95, 1.55], font_size=8.5)
    add_para(
        doc,
        "Commit graph evidence: full-repository main at 51e7a65; PantryPilot merge baseline 834e64e with tag v1.1.0; feature/safety-grocery-controls at 72a997e; original baseline 842ae5e with tag v1.0.0.",
        italic=True,
    )

    add_heading(doc, "Testing and Release Verification")
    add_para(
        doc,
        "Testing is part of the CM gate. The GitHub Actions workflow installs requirements.txt, runs pytest, and executes system_test_runner.py for PantryPilot changes on pushes and pull requests to main. Local verification produced 11 passing pytest tests and 6 passing system scenarios. The five new web tests exercise the demo page, allergen decision endpoint, grocery aggregation endpoint, pantry boundary endpoint, and quality-summary endpoint.",
    )
    add_para(
        doc,
        "The v1.1.0 tests cover release description, quality gates, allergen alias detection, safe recipe approval, grocery aggregation, and pantry-expiry boundary behavior. The added web/API regressions confirm that the interface delegates those decisions to the same deterministic backend controls. This evidence ties directly to unsafe-output, grocery-correctness, expiry-boundary, and UI/backend-consistency risks.",
    )

    doc.add_page_break()
    add_heading(doc, "Baseline, Release, and Rollback Control")
    add_caption(doc, "Table 5", "Baseline and release records")
    table = doc.add_table(rows=1, cols=5)
    for idx, header in enumerate(["Version", "Release name", "Tag", "Verification", "Rollback target"]):
        set_cell_text(table.rows[0].cells[idx], header, bold=True)
    rows = [
        ("1.0.0", "Plan & Adapt", "v1.0.0", "Initial source/config/test baseline", "N/A"),
        ("1.1.0", "Safety & Grocery Controls", "v1.1.0", "6 release scenarios passed; CI workflow configured", "v1.0.0"),
        ("Unreleased", "Web demonstration layer", "No tag", "11 pytest tests and 6 system scenarios passed locally", "v1.1.0"),
    ]
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
    apply_table_style(table, [0.8, 1.55, 0.85, 2.05, 1.1], font_size=8.8)
    rollback = add_para(
        doc,
        "Rollback is simple because tags identify recoverable baselines. If v1.1.0 fails release acceptance, the project can return to tag v1.0.0 while the feature branch is corrected and retested.",
    )
    rollback.paragraph_format.space_before = Pt(8)

    doc.add_page_break()
    add_heading(doc, "Repository Access and Source-Code Submission")
    add_para(
        doc,
        "The project is available in the full CISC-594 GitHub repository at github.com/liangzixuan/cisc-594. The PantryPilot directory, root-level workflow, branch history, annotated tags, tests, release notes, and CM documentation can therefore be inspected together. A separate source-code ZIP can be submitted through the LMS when the assignment requires file upload or an offline review copy.",
    )
    add_para(
        doc,
        "No secrets are stored in the submitted source. Environment-specific values belong in config/environment.example or in local environment variables, not in committed files. The .gitignore excludes pycache files, pytest cache, virtual environments, logs, and .env files.",
    )

    add_heading(doc, "Conclusion")
    add_para(
        doc,
        "PantryPilot's CM process satisfies the assignment by using version control, isolating baseline development on a feature branch, testing before merge, tagging completed versions, and documenting status and rollback records. The web demonstration adds new controlled items and test evidence without changing the meaning of the v1.1.0 tag. Keeping that work under [Unreleased] until it completes the same review and release gates makes the repository auditable, reproducible, and recoverable.",
    )

    add_heading(doc, "References")
    refs = [
        "Liang, Z. (2026a). PantryPilot configuration management generic files [Course assignment]. Harrisburg University of Science and Technology.",
        "Liang, Z. (2026b). PantryPilot risk management report [Course assignment]. Harrisburg University of Science and Technology.",
        "Liang, Z. (2026c). PantryPilot final project presentation [Course presentation]. Harrisburg University of Science and Technology.",
    ]
    for ref in refs:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.5)
        p.paragraph_format.space_after = Pt(6)
        p.add_run(ref)

    add_footer(doc)
    doc.save(OUT)


if __name__ == "__main__":
    build_report()
